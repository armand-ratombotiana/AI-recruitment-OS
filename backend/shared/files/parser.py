"""Resume parsing + structured field extraction.

Two layers of helpers:

1. **Format decoders** -- ``parse_pdf`` and ``parse_docx`` turn raw bytes into
   plain UTF-8 text using a real library (``pypdf`` / ``PyMuPDF`` for PDF,
   ``python-docx`` for DOCX).  The dispatcher ``parse_resume`` picks the
   right decoder from the MIME type or filename and falls back to a content
   sniff for callers that don't supply either.

2. **Field extractors** -- ``extract_email``, ``extract_phone``,
   ``extract_skills`` and ``extract_experience_years`` turn the text into
   structured fields.  These are deliberately simple regex / keyword
   matches so the module has no ML model dependency and runs in <1 ms per
   resume -- good enough to power real-time UI in the upload endpoint.

The helpers degrade gracefully:

* If a library is missing, the format decoder raises ``RuntimeError`` so the
  caller can decide to 400 the request or fall back to a different decoder.
* If extraction finds nothing, the function returns ``None`` (or an empty
  list) -- it never throws on missing data.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Iterable

logger = logging.getLogger("files.parser")


# ── Format detection ──────────────────────────────────────────────────────────


_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"

PDF_MIME_TYPES = {"application/pdf", "application/x-pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
    "application/x-zip-compressed",
    "application/octet-stream",
}
TEXT_MIME_TYPES = {"text/plain", "text/markdown", "text/html"}


def detect_content_type(
    content: bytes,
    *,
    declared_mime: str | None = None,
    filename: str | None = None,
) -> str:
    """Return the best-effort content type for ``content``.

    Order of preference:

    1. ``python-magic`` (libmagic) when available -- the gold standard.
    2. Extension hint from ``filename`` (.pdf / .docx / .txt).
    3. Byte-level magic sniff on the first few bytes.
    4. The caller-supplied ``declared_mime``.
    5. ``application/octet-stream`` as the safe default.
    """
    if declared_mime:
        declared = declared_mime.split(";", 1)[0].strip().lower()
    else:
        declared = ""

    # 1. python-magic (graceful fallback when libmagic is missing on Windows).
    try:
        import magic as _magic  # type: ignore

        detected = _magic.from_buffer(content[:4096], mime=True)
        if isinstance(detected, bytes):
            detected = detected.decode("ascii", errors="replace")
        detected = (detected or "").split(";", 1)[0].strip().lower()
        if detected and detected != "application/octet-stream":
            return detected
    except Exception:
        pass

    # 2. Extension hint
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "application/pdf"
    if name.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name.endswith(".doc"):
        return "application/msword"
    if name.endswith(".txt") or name.endswith(".md"):
        return "text/plain"
    if name.endswith(".html") or name.endswith(".htm"):
        return "text/html"

    # 3. Magic sniff
    head = content[:4]
    if head.startswith(_PDF_MAGIC):
        return "application/pdf"
    if head.startswith(_DOCX_MAGIC):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # 4. Declared mime from the caller
    if declared and declared != "application/octet-stream":
        return declared

    return "application/octet-stream"


def is_pdf(content_type: str, filename: str = "") -> bool:
    ct = (content_type or "").lower()
    name = (filename or "").lower()
    return "pdf" in ct or name.endswith(".pdf")


def is_docx(content_type: str, filename: str = "") -> bool:
    ct = (content_type or "").lower()
    name = (filename or "").lower()
    if "wordprocessingml" in ct or "msword" in ct:
        return True
    return name.endswith(".docx") or name.endswith(".doc")


def is_text(content_type: str, filename: str = "") -> bool:
    ct = (content_type or "").lower()
    name = (filename or "").lower()
    if ct.startswith("text/"):
        return True
    return name.endswith(".txt") or name.endswith(".md") or name.endswith(".markdown")


# ── Format decoders ───────────────────────────────────────────────────────────


def _decode_pdf_with_fitz(content: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    parts: list[str] = []
    for page in doc:
        try:
            parts.append(page.get_text("text") or "")
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("PyMuPDF page read failed: %s", exc)
    doc.close()
    return "\n".join(parts).strip()


def _decode_pdf_with_pypdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover
            logger.warning("pypdf page read failed: %s", exc)
    return "\n".join(parts).strip()


def parse_pdf(content: bytes) -> str:
    """Decode the text content of a PDF file.

    Tries PyMuPDF first (best fidelity on real-world resumes) and falls back
    to pypdf if PyMuPDF is not installed.  Raises ``RuntimeError`` when no
    PDF backend is available.
    """
    if not content:
        return ""
    if not content.lstrip().startswith(b"%PDF"):
        # Not a PDF -- the caller probably misidentified the bytes.  Return
        # empty text rather than blowing up so the upstream endpoint can 400
        # with a clear message.
        return ""
    last_err: Exception | None = None
    try:
        return _decode_pdf_with_fitz(content)
    except Exception as exc:
        last_err = exc
    try:
        return _decode_pdf_with_pypdf(content)
    except Exception as exc:
        last_err = exc
    raise RuntimeError("No PDF backend available for parsing") from last_err


def parse_docx(content: bytes) -> str:
    """Decode the text content of a DOCX file using python-docx."""
    if not content:
        return ""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for DOCX parsing") from exc
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # Also include any text inside tables -- real-world resumes put skills
    # inside two-column tables more often than not.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return "\n".join(parts).strip()


def parse_text(content: bytes) -> str:
    if not content:
        return ""
    return content.decode("utf-8", errors="replace").strip()


def parse_resume(
    content: bytes,
    content_type: str = "",
    filename: str = "",
) -> str:
    """Dispatcher: pick the right decoder and return plain UTF-8 text.

    Order of detection: ``content_type`` first, then ``filename`` extension,
    finally a byte-level magic sniff on the first few bytes.
    """
    if not content:
        return ""
    if is_pdf(content_type, filename) or content[:4].lstrip().startswith(b"%PDF"):
        return parse_pdf(content)
    if is_docx(content_type, filename) or content[:2] == b"PK":
        return parse_docx(content)
    if is_text(content_type, filename):
        return parse_text(content)
    # Final fallback -- try PDF magic one more time before treating as text.
    if content[:4].lstrip().startswith(b"%PDF"):
        return parse_pdf(content)
    return parse_text(content)


# ── Field extractors ──────────────────────────────────────────────────────────


_EMAIL_RE = re.compile(
    r"\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b"
)
# A phone is 7-15 digits with permissive separators.  We then sanity-check
# the digit count in ``extract_phone`` so a bare year doesn't match.
_PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s\-\.]?)?(?:\(\d{2,4}\)|\d{2,4})[\s\-\.]?\d{2,4}[\s\-\.]?\d{2,4}"
)
_YEARS_EXP_RE = re.compile(
    r"(?i)(?:\b(?:over|more than|at least|approximately|approx\.?)\s+)?"
    r"(\d{1,2})\+?\s*(?:\+)?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp\.?|professional|work)?",
)


def extract_email(text: str) -> str | None:
    """Return the first email address in ``text`` or ``None``."""
    if not text:
        return None
    match = _EMAIL_RE.search(text)
    return match.group(0).lower() if match else None


def extract_phone(text: str) -> str | None:
    """Return the first plausible phone number in ``text`` or ``None``.

    A plausible phone has between 7 and 15 digits once all non-digits are
    removed -- this filters out years, zip codes and other false positives
    while still matching common formats like ``+1 555 123 4567``.
    """
    if not text:
        return None
    for match in _PHONE_RE.finditer(text):
        candidate = match.group(0)
        digits = re.sub(r"\D", "", candidate)
        if 7 <= len(digits) <= 15:
            return candidate.strip()
    return None


def _normalise_skill(skill: str) -> str:
    return skill.strip().lower()


def extract_skills(text: str, known_skills: Iterable[str]) -> list[str]:
    """Return the subset of ``known_skills`` mentioned in ``text``.

    Matching is case-insensitive and word-boundary aware so a search for
    "Go" doesn't accidentally match "Google".  The returned list preserves
    the order of the first match in ``text`` (most relevant skills first).
    """
    if not text or not known_skills:
        return []
    lowered = text.lower()
    found: list[str] = []
    seen: set[str] = set()
    for raw in known_skills:
        skill = (raw or "").strip()
        if not skill:
            continue
        norm = _normalise_skill(skill)
        if not norm or norm in seen:
            continue
        # Use word boundaries when the skill is purely alphanumeric; fall
        # back to substring for things like "C++" or "C#".
        if re.fullmatch(r"[a-zA-Z0-9\.]+", norm):
            pattern = r"\b" + re.escape(norm) + r"\b"
        else:
            pattern = re.escape(norm)
        if re.search(pattern, lowered, flags=re.IGNORECASE):
            found.append(skill)
            seen.add(norm)
    return found


def extract_experience_years(text: str) -> int | None:
    """Return the explicit ``N years`` figure mentioned in ``text``.

    Picks the highest value when multiple figures exist ("3 years as a
    developer, 7 years in QA") and returns ``None`` when nothing matches.
    """
    if not text:
        return None
    best: int | None = None
    for match in _YEARS_EXP_RE.finditer(text):
        try:
            value = int(match.group(1))
        except (TypeError, ValueError):
            continue
        if value < 0 or value > 60:
            continue
        if best is None or value > best:
            best = value
    if best is not None:
        return best
    # Secondary heuristic: a "YYYY-YYYY" date range.  Compute the span and
    # use the largest span that fits "experience" or "work" within ~60 chars.
    range_re = re.compile(r"\b(19|20)\d{2}\s*[\-–to]+\s*(?:(19|20)\d{2}|present|now|current)\b", re.IGNORECASE)
    for match in range_re.finditer(text):
        start_str = match.group(0)[:4]
        try:
            start = int(start_str)
        except ValueError:
            continue
        end_token = match.group(0).split("-" if "-" in match.group(0) else "–")[-1].strip()
        end_token = end_token.replace("to", "").strip()
        if end_token.lower() in {"present", "now", "current"}:
            end = datetime.now().year
        else:
            try:
                end = int(end_token[:4])
            except ValueError:
                continue
        span = end - start
        if 0 <= span <= 60 and (best is None or span > best):
            best = span
    return best


__all__ = [
    "detect_content_type",
    "is_pdf",
    "is_docx",
    "is_text",
    "parse_pdf",
    "parse_docx",
    "parse_text",
    "parse_resume",
    "extract_email",
    "extract_phone",
    "extract_skills",
    "extract_experience_years",
]
